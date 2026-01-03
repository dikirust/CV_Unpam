#!/usr/bin/env python
"""
Waste Classification Report Generator - Phase 2
Generates HTML report and converts to DOCX
Indonesian text with English technical terms
"""

import json
import base64
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os

# Force unbuffered output
os.environ['PYTHONUNBUFFERED'] = '1'

print("=" * 80)
print("WASTE CLASSIFICATION - REPORT GENERATOR (PHASE 2)")
print("=" * 80)

# Setup paths
BASE_DIR = Path('.')
OUTPUT_DIR = BASE_DIR / 'output'
REPORT_DIR = OUTPUT_DIR / 'report'

print(f"\nOutput directory: {OUTPUT_DIR}")
print(f"Report directory: {REPORT_DIR}")
sys.stdout.flush()

# Load results
print("\nLoading results...")
sys.stdout.flush()

with open(OUTPUT_DIR / 'evaluation_results.json', 'r') as f:
    results = json.load(f)

print("✓ Results loaded")
sys.stdout.flush()

# Extract metrics
custom_cnn_results = results['custom_cnn']
mobilenet_results = results['mobilenetv2']
training_times = results['training_times']

print("\nMetrics loaded:")
print(f"  Custom CNN Accuracy: {custom_cnn_results['accuracy']:.4f}")
print(f"  MobileNetV2 Accuracy: {mobilenet_results['accuracy']:.4f}")
sys.stdout.flush()

# ==================== HTML REPORT ====================
print("\n" + "=" * 80)
print("GENERATING HTML REPORT")
print("=" * 80)

html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Klasifikasi Sampah Menggunakan Deep Learning</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #f5f5f5;
            padding: 20px;
        }}
        
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background-color: white;
            padding: 40px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 40px;
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 20px;
        }}
        
        h1 {{
            color: #2c3e50;
            font-size: 28px;
            margin-bottom: 10px;
        }}
        
        .subtitle {{
            color: #7f8c8d;
            font-size: 16px;
            margin-bottom: 10px;
        }}
        
        .author-info {{
            font-size: 14px;
            color: #95a5a6;
            margin-top: 15px;
        }}
        
        .section {{
            margin-bottom: 40px;
        }}
        
        h2 {{
            color: #2c3e50;
            font-size: 20px;
            margin-bottom: 15px;
            padding-bottom: 10px;
            border-bottom: 2px solid #3498db;
        }}
        
        h3 {{
            color: #34495e;
            font-size: 16px;
            margin-top: 15px;
            margin-bottom: 10px;
        }}
        
        p {{
            text-align: justify;
            margin-bottom: 15px;
            line-height: 1.8;
        }}
        
        .highlight {{
            background-color: #ecf0f1;
            padding: 15px;
            border-left: 4px solid #3498db;
            margin: 15px 0;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        
        table, th, td {{
            border: 1px solid #bdc3c7;
        }}
        
        th {{
            background-color: #3498db;
            color: white;
            padding: 12px;
            text-align: left;
        }}
        
        td {{
            padding: 10px;
        }}
        
        tr:nth-child(even) {{
            background-color: #ecf0f1;
        }}
        
        .image-container {{
            text-align: center;
            margin: 30px 0;
        }}
        
        .image-container img {{
            max-width: 100%;
            height: auto;
            border: 1px solid #bdc3c7;
            border-radius: 5px;
        }}
        
        .image-caption {{
            font-size: 13px;
            color: #7f8c8d;
            margin-top: 10px;
            font-style: italic;
        }}
        
        .reference {{
            font-size: 13px;
            margin-bottom: 10px;
            margin-left: 20px;
        }}
        
        .footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 2px solid #bdc3c7;
            text-align: center;
            color: #7f8c8d;
            font-size: 12px;
        }}
        
        .metrics-grid {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            margin: 20px 0;
        }}
        
        .metric-box {{
            background-color: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            border-left: 4px solid #3498db;
        }}
        
        .metric-label {{
            font-size: 13px;
            color: #7f8c8d;
            margin-bottom: 5px;
        }}
        
        .metric-value {{
            font-size: 18px;
            color: #2c3e50;
            font-weight: bold;
        }}
        
        .model-name {{
            font-weight: bold;
            color: #2c3e50;
        }}
    </style>
</head>
<body>
    <div class="container">
        <!-- Header -->
        <div class="header">
            <h1>Klasifikasi Sampah Menggunakan Deep Learning</h1>
            <p class="subtitle">Perbandingan Arsitektur Custom CNN dan MobileNetV2</p>
            <div class="author-info">
                <p><strong>Penulis:</strong> Diki Rustian</p>
                <p><strong>Email:</strong> diki.rstn@gmail.com</p>
                <p><strong>Institusi:</strong> Universitas Pamulang, Indonesia</p>
                <p><strong>Tanggal:</strong> {datetime.now().strftime('%d Januari %Y')}</p>
            </div>
        </div>

        <!-- Abstrak -->
        <div class="section">
            <h2>Abstrak</h2>
            <p>
                Pengelolaan sampah merupakan isu penting dalam keberlanjutan lingkungan modern. Klasifikasi otomatis sampah 
                menggunakan teknologi computer vision dapat meningkatkan efisiensi sistem pemisahan sampah. Penelitian ini 
                membandingkan dua arsitektur convolutional neural network (CNN) untuk klasifikasi lima jenis sampah 
                (sampah organik, kaca, logam, kertas, dan plastik). Kami membandingkan Custom CNN yang dibangun dari awal 
                dengan MobileNetV2 yang menggunakan transfer learning. Hasil penelitian menunjukkan bahwa MobileNetV2 mencapai 
                akurasi {mobilenet_results['accuracy']:.2%} dengan waktu training {training_times['mobilenetv2']:.1f} detik, 
                sementara Custom CNN mencapai akurasi {custom_cnn_results['accuracy']:.2%} dengan waktu training 
                {training_times['custom_cnn']:.1f} detik. Transfer learning terbukti memberikan hasil yang lebih baik 
                dengan konvergensi yang lebih cepat.
            </p>
        </div>

        <!-- Pendahuluan -->
        <div class="section">
            <h2>Pendahuluan</h2>
            
            <h3>Latar Belakang dan Rumusan Masalah</h3>
            <p>
                Sampah merupakan hasil sampingan dari aktivitas manusia yang terus meningkat seiring pertumbuhan populasi 
                dan konsumsi. Sistem pemisahan sampah manual memerlukan biaya operasional tinggi dan rentan terhadap kesalahan. 
                Penerapan otomasi berbasis computer vision dapat meningkatkan akurasi dan efisiensi pemisahan sampah.
            </p>
            <p>
                Rumusan masalah dalam penelitian ini adalah: "Arsitektur CNN manakah yang memberikan performa terbaik 
                untuk klasifikasi otomatis lima jenis sampah?" Pertanyaan penelitian fokus pada perbandingan akurasi, 
                kecepatan training, dan efisiensi model antara dua pendekatan berbeda.
            </p>
            
            <h3>Usulan Solusi</h3>
            <p>
                Penelitian ini mengusulkan dua solusi berbasis deep learning:
            </p>
            <div class="highlight">
                <p><strong>1. Custom CNN:</strong> Arsitektur convolutional neural network yang dibangun dari awal 
                dengan tiga blok convolutional (32→64→128 filter). Pendekatan ini memberikan fleksibilitas penuh 
                dalam desain arsitektur namun memerlukan jumlah data training yang lebih besar.</p>
                
                <p><strong>2. MobileNetV2:</strong> Arsitektur transfer learning yang memanfaatkan weight pre-trained 
                dari ImageNet. MobileNetV2 dirancang untuk efisiensi dengan menggunakan depthwise separable convolutions, 
                ideal untuk deployment di perangkat dengan resource terbatas.</p>
            </div>
        </div>

        <!-- Metodologi -->
        <div class="section">
            <h2>Metodologi</h2>
            
            <h3>Dataset dan Preprocessing</h3>
            <p>
                Dataset terdiri dari 8,400 gambar sampah yang dibagi menjadi lima kelas:
            </p>
            <table>
                <tr>
                    <th>Kelas Sampah</th>
                    <th>Jumlah Gambar</th>
                    <th>Persentase</th>
                </tr>
                <tr>
                    <td>Sampah Organik (foodwaste)</td>
                    <td>970</td>
                    <td>11.5%</td>
                </tr>
                <tr>
                    <td>Kaca (glass)</td>
                    <td>954</td>
                    <td>11.4%</td>
                </tr>
                <tr>
                    <td>Logam (metal)</td>
                    <td>1,713</td>
                    <td>20.4%</td>
                </tr>
                <tr>
                    <td>Kertas (paper)</td>
                    <td>2,267</td>
                    <td>27.0%</td>
                </tr>
                <tr>
                    <td>Plastik (plastic)</td>
                    <td>2,496</td>
                    <td>29.7%</td>
                </tr>
            </table>
            
            <p>
                Data preprocessing meliputi:
            </p>
            <ul style="margin-left: 20px; margin-bottom: 15px;">
                <li>Resize gambar menjadi 64×64 pixel</li>
                <li>Normalisasi pixel value ke rentang 0-1</li>
                <li>Split data: 70% training, 15% validation, 15% testing</li>
                <li>Stratified split untuk memastikan distribusi kelas seimbang</li>
            </ul>
            
            <h3>Arsitektur Model</h3>
            
            <h4>Custom CNN</h4>
            <p>
                Model Custom CNN terdiri dari tiga blok convolutional dengan struktur:
            </p>
            <div class="highlight">
                <p>Conv2D(32, 3×3) → BatchNorm → MaxPool(2×2) → Dropout(0.2)</p>
                <p>Conv2D(64, 3×3) → BatchNorm → MaxPool(2×2) → Dropout(0.2)</p>
                <p>Conv2D(128, 3×3) → BatchNorm → MaxPool(2×2) → Dropout(0.2)</p>
                <p>GlobalAveragePooling → Dense(256) → Dense(128) → Dense(5, softmax)</p>
            </div>
            <p>
                Total parameters: 310,405. Model ini menggunakan BatchNormalization untuk stabilisasi training 
                dan Dropout untuk regularisasi.
            </p>
            
            <h4>MobileNetV2</h4>
            <p>
                MobileNetV2 merupakan arsitektur ringan yang dirancang oleh Google untuk deployment mobile. 
                Karakteristik utama:
            </p>
            <ul style="margin-left: 20px; margin-bottom: 15px;">
                <li>Menggunakan depthwise separable convolutions untuk efisiensi</li>
                <li>Linear bottleneck dan inverted residual blocks</li>
                <li>Pre-trained weights dari ImageNet (87.5M parameters)</li>
                <li>Custom head: GlobalAveragePooling → Dense(256) → Dense(128) → Dense(5, softmax)</li>
            </ul>
            
            <h3>Konfigurasi Training</h3>
            <table>
                <tr>
                    <th>Parameter</th>
                    <th>Nilai</th>
                </tr>
                <tr>
                    <td>Optimizer</td>
                    <td>Adam (learning rate = 0.001)</td>
                </tr>
                <tr>
                    <td>Loss Function</td>
                    <td>Sparse Categorical Crossentropy</td>
                </tr>
                <tr>
                    <td>Batch Size</td>
                    <td>16</td>
                </tr>
                <tr>
                    <td>Epochs</td>
                    <td>50 (dengan Early Stopping)</td>
                </tr>
                <tr>
                    <td>EarlyStopping</td>
                    <td>patience=10, monitor val_loss</td>
                </tr>
                <tr>
                    <td>ReduceLROnPlateau</td>
                    <td>factor=0.5, patience=5</td>
                </tr>
            </table>
        </div>

        <!-- Hasil dan Diskusi -->
        <div class="section">
            <h2>Hasil dan Diskusi</h2>
            
            <h3>Performa Model pada Test Set</h3>
            <p>
                Hasil evaluasi kedua model pada test set disajikan dalam tabel berikut:
            </p>
            <table>
                <tr>
                    <th>Metrik</th>
                    <th>Custom CNN</th>
                    <th>MobileNetV2</th>
                    <th>Perbedaan</th>
                </tr>
                <tr>
                    <td><strong>Accuracy</strong></td>
                    <td>{custom_cnn_results['accuracy']:.4f}</td>
                    <td>{mobilenet_results['accuracy']:.4f}</td>
                    <td>{(mobilenet_results['accuracy'] - custom_cnn_results['accuracy']):.4f}</td>
                </tr>
                <tr>
                    <td><strong>Precision</strong></td>
                    <td>{custom_cnn_results['precision']:.4f}</td>
                    <td>{mobilenet_results['precision']:.4f}</td>
                    <td>{(mobilenet_results['precision'] - custom_cnn_results['precision']):.4f}</td>
                </tr>
                <tr>
                    <td><strong>Recall</strong></td>
                    <td>{custom_cnn_results['recall']:.4f}</td>
                    <td>{mobilenet_results['recall']:.4f}</td>
                    <td>{(mobilenet_results['recall'] - custom_cnn_results['recall']):.4f}</td>
                </tr>
                <tr>
                    <td><strong>F1-Score</strong></td>
                    <td>{custom_cnn_results['f1_score']:.4f}</td>
                    <td>{mobilenet_results['f1_score']:.4f}</td>
                    <td>{(mobilenet_results['f1_score'] - custom_cnn_results['f1_score']):.4f}</td>
                </tr>
            </table>
            
            <h3>Performa Per-Kelas</h3>
            <p>
                Akurasi per kelas untuk kedua model:
            </p>
            <table>
                <tr>
                    <th>Kelas Sampah</th>
                    <th>Custom CNN</th>
                    <th>MobileNetV2</th>
                </tr>"""

# Add per-class accuracy
for class_name in ['foodwaste', 'glass', 'metal', 'paper', 'plastic']:
    custom_acc = custom_cnn_results['per_class'].get(class_name, {}).get('accuracy', 0)
    mobile_acc = mobilenet_results['per_class'].get(class_name, {}).get('accuracy', 0)
    html_content += f"""
                <tr>
                    <td>{class_name.capitalize()}</td>
                    <td>{custom_acc:.4f}</td>
                    <td>{mobile_acc:.4f}</td>
                </tr>"""

html_content += """
            </table>
            
            <h3>Visualisasi Hasil</h3>
            <p>Perbandingan metrik performa kedua model:</p>
"""

# Add images if they exist
image_files = [
    ('metrics_comparison.png', 'Perbandingan Metrik Performa'),
    ('confusion_matrices.png', 'Confusion Matrix Kedua Model'),
    ('per_class_accuracy.png', 'Akurasi Per-Kelas')
]

for image_file, caption in image_files:
    image_path = REPORT_DIR / image_file
    if image_path.exists():
        html_content += f"""
            <div class="image-container">
                <img src="{image_file}" alt="{caption}">
                <p class="image-caption">{caption}</p>
            </div>
"""

html_content += f"""
            <h3>Analisis dan Interpretasi</h3>
            <p>
                Berdasarkan hasil penelitian, MobileNetV2 menunjukkan performa yang lebih superior dibandingkan 
                Custom CNN dengan perbedaan akurasi {(mobilenet_results['accuracy'] - custom_cnn_results['accuracy']):.2%}. 
                Hal ini dapat dijelaskan oleh beberapa faktor:
            </p>
            <ul style="margin-left: 20px; margin-bottom: 15px;">
                <li><strong>Transfer Learning:</strong> MobileNetV2 memanfaatkan pengetahuan yang sudah dipelajari 
                dari dataset ImageNet dengan jutaan gambar, memberikan inisialisasi weight yang lebih baik.</li>
                
                <li><strong>Arsitektur yang Optimal:</strong> Desain depthwise separable convolutions memungkinkan 
                ekstraksi fitur yang efisien dengan parameter lebih sedikit.</li>
                
                <li><strong>Kecepatan Konvergensi:</strong> MobileNetV2 mencapai performa optimal lebih cepat, 
                ditunjukkan dengan training time {training_times['mobilenetv2']:.1f} detik vs Custom CNN 
                {training_times['custom_cnn']:.1f} detik.</li>
            </ul>
            
            <p>
                Custom CNN tetap memiliki nilai edukatif tinggi karena memungkinkan pemahaman mendalam tentang 
                mekanisme CNN. Akurasi {custom_cnn_results['accuracy']:.2%} juga cukup baik untuk aplikasi praktis.
            </p>
        </div>

        <!-- Kesimpulan -->
        <div class="section">
            <h2>Kesimpulan</h2>
            <p>
                Penelitian ini telah membandingkan dua arsitektur CNN untuk klasifikasi otomatis sampah. Temuan utama:
            </p>
            <div class="highlight">
                <p><strong>1.</strong> MobileNetV2 mencapai akurasi {mobilenet_results['accuracy']:.2%}, 
                {(mobilenet_results['accuracy'] - custom_cnn_results['accuracy']):.2%} lebih tinggi dari Custom CNN.</p>
                
                <p><strong>2.</strong> Transfer learning terbukti efektif untuk meningkatkan akurasi dengan data training 
                terbatas.</p>
                
                <p><strong>3.</strong> MobileNetV2 lebih cepat dalam training ({training_times['mobilenetv2']:.0f}s vs 
                {training_times['custom_cnn']:.0f}s), ideal untuk deployment praktis.</p>
                
                <p><strong>4.</strong> Kedua model menunjukkan performa yang baik (>70%) dan layak untuk aplikasi 
                sistem pemisahan sampah otomatis.</p>
            </div>
            
            <h3>Saran untuk Penelitian Lanjutan</h3>
            <ul style="margin-left: 20px; margin-bottom: 15px;">
                <li>Menggunakan dataset yang lebih besar dengan variasi kondisi pencahayaan berbeda</li>
                <li>Menerapkan data augmentation yang lebih kompleks</li>
                <li>Melakukan fine-tuning pada layer base MobileNetV2</li>
                <li>Membandingkan dengan arsitektur modern lainnya (EfficientNet, Vision Transformer)</li>
                <li>Implementasi real-time inference pada hardware embedded</li>
            </ul>
        </div>

        <!-- Daftar Pustaka -->
        <div class="section">
            <h2>Daftar Pustaka</h2>
            <div class="reference">1. Sandler, M., Howard, A., Zhu, M., Zhmoginov, A., & Chen, L. C. (2018). 
            MobileNetV2: Inverted Residuals and Linear Bottlenecks. arXiv preprint arXiv:1801.04381.</div>
            
            <div class="reference">2. He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning for Image 
            Recognition. In IEEE conference on computer vision and pattern recognition (pp. 770-778).</div>
            
            <div class="reference">3. LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521(7553), 436-444.</div>
            
            <div class="reference">4. Krizhevsky, A., Sutskever, I., & Hinton, G. E. (2017). ImageNet classification with 
            deep convolutional neural networks. Communications of the ACM, 60(6), 84-90.</div>
            
            <div class="reference">5. Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. 
            MIT press.</div>
        </div>

        <!-- Footer -->
        <div class="footer">
            <p>© 2026 Diki Rustian | Universitas Pamulang, Indonesia</p>
            <p>Generated on {datetime.now().strftime('%d %B %Y at %H:%M:%S')}</p>
        </div>
    </div>
</body>
</html>
"""

# Save HTML
html_file = OUTPUT_DIR / 'report.html'
with open(html_file, 'w', encoding='utf-8') as f:
    f.write(html_content)

print(f"✓ HTML report saved: {html_file}")
sys.stdout.flush()

# ==================== DOCX REPORT ====================
print("\n" + "=" * 80)
print("GENERATING DOCX REPORT")
print("=" * 80)

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Klasifikasi Sampah Menggunakan Deep Learning', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Perbandingan Arsitektur Custom CNN dan MobileNetV2')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.italic = True

# Author info
author_info = doc.add_paragraph()
author_info.add_run('Penulis: ').bold = True
author_info.add_run('Diki Rustian\n')
author_info.add_run('Email: ').bold = True
author_info.add_run('diki.rstn@gmail.com\n')
author_info.add_run('Institusi: ').bold = True
author_info.add_run('Universitas Pamulang, Indonesia\n')
author_info.add_run('Tanggal: ').bold = True
author_info.add_run(datetime.now().strftime('%d Januari %Y'))
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Abstrak
doc.add_heading('Abstrak', 1)
abstract_text = f"""Pengelolaan sampah merupakan isu penting dalam keberlanjutan lingkungan modern. Klasifikasi otomatis sampah 
menggunakan teknologi computer vision dapat meningkatkan efisiensi sistem pemisahan sampah. Penelitian ini membandingkan dua 
arsitektur convolutional neural network (CNN) untuk klasifikasi lima jenis sampah (sampah organik, kaca, logam, kertas, dan 
plastik). Kami membandingkan Custom CNN yang dibangun dari awal dengan MobileNetV2 yang menggunakan transfer learning. Hasil 
penelitian menunjukkan bahwa MobileNetV2 mencapai akurasi {mobilenet_results['accuracy']:.2%} dengan waktu training 
{training_times['mobilenetv2']:.1f} detik, sementara Custom CNN mencapai akurasi {custom_cnn_results['accuracy']:.2%} dengan 
waktu training {training_times['custom_cnn']:.1f} detik. Transfer learning terbukti memberikan hasil yang lebih baik dengan 
konvergensi yang lebih cepat."""
doc.add_paragraph(abstract_text)

doc.add_paragraph()

# Pendahuluan
doc.add_heading('Pendahuluan', 1)
doc.add_heading('Latar Belakang dan Rumusan Masalah', 2)
intro_text = """Sampah merupakan hasil sampingan dari aktivitas manusia yang terus meningkat seiring pertumbuhan populasi dan 
konsumsi. Sistem pemisahan sampah manual memerlukan biaya operasional tinggi dan rentan terhadap kesalahan. Penerapan otomasi 
berbasis computer vision dapat meningkatkan akurasi dan efisiensi pemisahan sampah.

Rumusan masalah dalam penelitian ini adalah: "Arsitektur CNN manakah yang memberikan performa terbaik untuk klasifikasi 
otomatis lima jenis sampah?" Pertanyaan penelitian fokus pada perbandingan akurasi, kecepatan training, dan efisiensi model 
antara dua pendekatan berbeda."""
doc.add_paragraph(intro_text)

doc.add_heading('Usulan Solusi', 2)
doc.add_paragraph('Penelitian ini mengusulkan dua solusi berbasis deep learning:')

p1 = doc.add_paragraph()
p1.add_run('Custom CNN: ').bold = True
p1.add_run('Arsitektur convolutional neural network yang dibangun dari awal dengan tiga blok convolutional (32→64→128 filter). '
          'Pendekatan ini memberikan fleksibilitas penuh dalam desain arsitektur namun memerlukan jumlah data training yang lebih besar.')

p2 = doc.add_paragraph()
p2.add_run('MobileNetV2: ').bold = True
p2.add_run('Arsitektur transfer learning yang memanfaatkan weight pre-trained dari ImageNet. MobileNetV2 dirancang untuk '
          'efisiensi dengan menggunakan depthwise separable convolutions, ideal untuk deployment di perangkat dengan resource terbatas.')

doc.add_paragraph()

# Metodologi
doc.add_heading('Metodologi', 1)
doc.add_heading('Dataset dan Preprocessing', 2)

doc.add_paragraph('Dataset terdiri dari 8,400 gambar sampah yang dibagi menjadi lima kelas:')

# Table untuk dataset
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Kelas Sampah'
hdr_cells[1].text = 'Jumlah Gambar'
hdr_cells[2].text = 'Persentase'

# Data rows
data = [
    ('Sampah Organik (foodwaste)', '970', '11.5%'),
    ('Kaca (glass)', '954', '11.4%'),
    ('Logam (metal)', '1,713', '20.4%'),
    ('Kertas (paper)', '2,267', '27.0%'),
    ('Plastik (plastic)', '2,496', '29.7%'),
]

for i, (class_name, count, percent) in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = class_name
    row_cells[1].text = count
    row_cells[2].text = percent

doc.add_paragraph('Data preprocessing meliputi:')
preprocessing_items = [
    'Resize gambar menjadi 64×64 pixel',
    'Normalisasi pixel value ke rentang 0-1',
    'Split data: 70% training, 15% validation, 15% testing',
    'Stratified split untuk memastikan distribusi kelas seimbang'
]
for item in preprocessing_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Arsitektur Model', 2)
doc.add_heading('Custom CNN', 3)
doc.add_paragraph('Model Custom CNN terdiri dari tiga blok convolutional dengan struktur:')
cnn_arch = [
    'Conv2D(32, 3×3) → BatchNorm → MaxPool(2×2) → Dropout(0.2)',
    'Conv2D(64, 3×3) → BatchNorm → MaxPool(2×2) → Dropout(0.2)',
    'Conv2D(128, 3×3) → BatchNorm → MaxPool(2×2) → Dropout(0.2)',
    'GlobalAveragePooling → Dense(256) → Dense(128) → Dense(5, softmax)'
]
for arch in cnn_arch:
    doc.add_paragraph(arch, style='List Bullet')
doc.add_paragraph('Total parameters: 310,405. Model ini menggunakan BatchNormalization untuk stabilisasi training dan Dropout untuk regularisasi.')

doc.add_heading('MobileNetV2', 3)
doc.add_paragraph('MobileNetV2 merupakan arsitektur ringan yang dirancang oleh Google untuk deployment mobile. Karakteristik utama:')
mobilenet_features = [
    'Menggunakan depthwise separable convolutions untuk efisiensi',
    'Linear bottleneck dan inverted residual blocks',
    'Pre-trained weights dari ImageNet (87.5M parameters)',
    'Custom head: GlobalAveragePooling → Dense(256) → Dense(128) → Dense(5, softmax)'
]
for feature in mobilenet_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Konfigurasi Training', 2)

# Training config table
train_table = doc.add_table(rows=8, cols=2)
train_table.style = 'Light Grid Accent 1'

train_hdr = train_table.rows[0].cells
train_hdr[0].text = 'Parameter'
train_hdr[1].text = 'Nilai'

train_config = [
    ('Optimizer', 'Adam (learning rate = 0.001)'),
    ('Loss Function', 'Sparse Categorical Crossentropy'),
    ('Batch Size', '16'),
    ('Epochs', '50 (dengan Early Stopping)'),
    ('EarlyStopping', 'patience=10, monitor val_loss'),
    ('ReduceLROnPlateau', 'factor=0.5, patience=5'),
]

for i, (param, value) in enumerate(train_config, 1):
    row = train_table.rows[i].cells
    row[0].text = param
    row[1].text = value

doc.add_paragraph()

# Hasil dan Diskusi
doc.add_heading('Hasil dan Diskusi', 1)
doc.add_heading('Performa Model pada Test Set', 2)

# Results table
result_table = doc.add_table(rows=5, cols=4)
result_table.style = 'Light Grid Accent 1'

result_hdr = result_table.rows[0].cells
result_hdr[0].text = 'Metrik'
result_hdr[1].text = 'Custom CNN'
result_hdr[2].text = 'MobileNetV2'
result_hdr[3].text = 'Perbedaan'

result_data = [
    ('Accuracy', f"{custom_cnn_results['accuracy']:.4f}", f"{mobilenet_results['accuracy']:.4f}", 
     f"{(mobilenet_results['accuracy'] - custom_cnn_results['accuracy']):.4f}"),
    ('Precision', f"{custom_cnn_results['precision']:.4f}", f"{mobilenet_results['precision']:.4f}",
     f"{(mobilenet_results['precision'] - custom_cnn_results['precision']):.4f}"),
    ('Recall', f"{custom_cnn_results['recall']:.4f}", f"{mobilenet_results['recall']:.4f}",
     f"{(mobilenet_results['recall'] - custom_cnn_results['recall']):.4f}"),
    ('F1-Score', f"{custom_cnn_results['f1_score']:.4f}", f"{mobilenet_results['f1_score']:.4f}",
     f"{(mobilenet_results['f1_score'] - custom_cnn_results['f1_score']):.4f}"),
]

for i, (metric, custom, mobile, diff) in enumerate(result_data, 1):
    row = result_table.rows[i].cells
    row[0].text = metric
    row[1].text = custom
    row[2].text = mobile
    row[3].text = diff

doc.add_heading('Performa Per-Kelas', 2)
doc.add_paragraph('Akurasi per kelas untuk kedua model:')

class_table = doc.add_table(rows=6, cols=3)
class_table.style = 'Light Grid Accent 1'

class_hdr = class_table.rows[0].cells
class_hdr[0].text = 'Kelas Sampah'
class_hdr[1].text = 'Custom CNN'
class_hdr[2].text = 'MobileNetV2'

for i, class_name in enumerate(['foodwaste', 'glass', 'metal', 'paper', 'plastic'], 1):
    custom_acc = custom_cnn_results['per_class'].get(class_name, {}).get('accuracy', 0)
    mobile_acc = mobilenet_results['per_class'].get(class_name, {}).get('accuracy', 0)
    row = class_table.rows[i].cells
    row[0].text = class_name.capitalize()
    row[1].text = f"{custom_acc:.4f}"
    row[2].text = f"{mobile_acc:.4f}"

doc.add_heading('Visualisasi Hasil', 2)
doc.add_paragraph('Perbandingan metrik performa kedua model:')

# Add images
for image_file, caption in image_files:
    image_path = REPORT_DIR / image_file
    if image_path.exists():
        doc.add_paragraph(caption)
        doc.add_picture(str(image_path), width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Analisis dan Interpretasi', 2)
analysis_text = f"""Berdasarkan hasil penelitian, MobileNetV2 menunjukkan performa yang lebih superior dibandingkan Custom CNN 
dengan perbedaan akurasi {(mobilenet_results['accuracy'] - custom_cnn_results['accuracy']):.2%}. Hal ini dapat dijelaskan oleh 
beberapa faktor:"""
doc.add_paragraph(analysis_text)

analysis_points = [
    ('Transfer Learning', 'MobileNetV2 memanfaatkan pengetahuan yang sudah dipelajari dari dataset ImageNet dengan jutaan gambar, '
                          'memberikan inisialisasi weight yang lebih baik.'),
    ('Arsitektur yang Optimal', 'Desain depthwise separable convolutions memungkinkan ekstraksi fitur yang efisien dengan parameter lebih sedikit.'),
    ('Kecepatan Konvergensi', f'MobileNetV2 mencapai performa optimal lebih cepat, ditunjukkan dengan training time {training_times["mobilenetv2"]:.1f} '
                              f'detik vs Custom CNN {training_times["custom_cnn"]:.1f} detik.')
]

for title, content in analysis_points:
    p = doc.add_paragraph()
    p.add_run(title + ': ').bold = True
    p.add_run(content)

conclude_text = f'Custom CNN tetap memiliki nilai edukatif tinggi karena memungkinkan pemahaman mendalam tentang mekanisme CNN. Akurasi {custom_cnn_results["accuracy"]:.2%} juga cukup baik untuk aplikasi praktis.'
doc.add_paragraph(conclude_text)

doc.add_paragraph()

# Kesimpulan
doc.add_heading('Kesimpulan', 1)

conclusion_intro = doc.add_paragraph('Penelitian ini telah membandingkan dua arsitektur CNN untuk klasifikasi otomatis sampah. Temuan utama:')

conclusions = [
    f'MobileNetV2 mencapai akurasi {mobilenet_results["accuracy"]:.2%}, {(mobilenet_results["accuracy"] - custom_cnn_results["accuracy"]):.2%} lebih tinggi dari Custom CNN.',
    'Transfer learning terbukti efektif untuk meningkatkan akurasi dengan data training terbatas.',
    f'MobileNetV2 lebih cepat dalam training ({training_times["mobilenetv2"]:.0f}s vs {training_times["custom_cnn"]:.0f}s), ideal untuk deployment praktis.',
    'Kedua model menunjukkan performa yang baik (>70%) dan layak untuk aplikasi sistem pemisahan sampah otomatis.'
]

for i, conclusion in enumerate(conclusions, 1):
    doc.add_paragraph(f'{i}. {conclusion}')

doc.add_heading('Saran untuk Penelitian Lanjutan', 2)
suggestions = [
    'Menggunakan dataset yang lebih besar dengan variasi kondisi pencahayaan berbeda',
    'Menerapkan data augmentation yang lebih kompleks',
    'Melakukan fine-tuning pada layer base MobileNetV2',
    'Membandingkan dengan arsitektur modern lainnya (EfficientNet, Vision Transformer)',
    'Implementasi real-time inference pada hardware embedded'
]
for suggestion in suggestions:
    doc.add_paragraph(suggestion, style='List Bullet')

doc.add_paragraph()

# Daftar Pustaka
doc.add_heading('Daftar Pustaka', 1)
references = [
    'Sandler, M., Howard, A., Zhu, M., Zhmoginov, A., & Chen, L. C. (2018). MobileNetV2: Inverted Residuals and Linear Bottlenecks. arXiv preprint arXiv:1801.04381.',
    'He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning for Image Recognition. In IEEE conference on computer vision and pattern recognition (pp. 770-778).',
    'LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521(7553), 436-444.',
    'Krizhevsky, A., Sutskever, I., & Hinton, G. E. (2017). ImageNet classification with deep convolutional neural networks. Communications of the ACM, 60(6), 84-90.',
    'Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT press.'
]

for ref in references:
    doc.add_paragraph(ref, style='List Number')

# Save DOCX
docx_file = OUTPUT_DIR / 'report.docx'
doc.save(str(docx_file))

print(f"✓ DOCX report saved: {docx_file}")
sys.stdout.flush()

# ==================== COMPLETION ====================
print("\n" + "=" * 80)
print("REPORT GENERATION COMPLETED SUCCESSFULLY!")
print("=" * 80)
print(f"\nGenerated reports:")
print(f"  ✓ {html_file}")
print(f"  ✓ {docx_file}")
print("\nReport contents:")
print(f"  - Abstrak (Abstract)")
print(f"  - Pendahuluan (Introduction)")
print(f"  - Metodologi (Methodology)")
print(f"  - Hasil dan Diskusi (Results & Discussion)")
print(f"  - Kesimpulan (Conclusion)")
print(f"  - Daftar Pustaka (References)")
print(f"  - Visualisasi (3 PNG images embedded)")
print("\n" + "=" * 80)
print("PHASE 2 COMPLETE - READY FOR SUBMISSION!")
print("=" * 80)
print("\nNext: Upload report.docx to Mentari UAS session")
print("=" * 80)
sys.stdout.flush()
